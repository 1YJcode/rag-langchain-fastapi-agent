import hashlib
import os, asyncio, aiofiles
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredMarkdownLoader, \
    UnstructuredPowerPointLoader
from langchain_core.documents import Document
from backend.app.core.logger_handler import logger
from backend.app.utils.path_tool import get_abstract_path


async def get_file_md5_hex(file_path: str) -> str:
    """
    获取文件的md5值
    :return:
    """
    # 处理路径，确保使用绝对路径
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path

    if not os.path.exists(abs_file_path):
        logger.error(f"[md5计算]文件路径{abs_file_path}不存在！")
        return ""

    if not os.path.isfile(abs_file_path):
        logger.error(f"[md5计算]文件路径{abs_file_path}不是文件！")
        return ""

    # 获得hashlib.md5对象
    md5_obj = hashlib.md5()
    # 1KB分片，避免文件过大爆内存
    chunk_size = 1024
    try:
        async with aiofiles.open(abs_file_path, 'rb') as f:  # 以分片读取，必须用二进制读取模式
            while chunk := await f.read(chunk_size):
                md5_obj.update(chunk)

            """
            chunk = f.read(chunk_size)
            while chunk:
                md5_obj.update(chunk)
                chunk = f.read(chunk_size)
            """
            md5_hex = md5_obj.hexdigest()
            return md5_hex

    except Exception as e:
        logger.error(f"[md5计算]读取文件{abs_file_path}时出错:{e}")
        return ""


async def listdir_allowed_type(path: str, allowed_type: tuple[str]) -> tuple:
    """
    获取目录下所有允许的文件类型
    :return:
    """
    # 处理路径，确保使用绝对路径
    abs_file_path = get_abstract_path(path) if not os.path.isabs(path) else path

    if not os.path.exists(abs_file_path):
        logger.error(f"[文件列表]目录路径{abs_file_path}不存在！")
        return ()
    if not os.path.isdir(abs_file_path):
        logger.error(f"[文件列表]目录路径{abs_file_path}不是文件夹！")
        return ()

    file_path = []
    for f in await asyncio.to_thread(os.listdir, abs_file_path):
        if f.endswith(allowed_type):
            os.path.join(abs_file_path, f)
            file_path.append(f)

    return tuple(file_path)


async def pdf_loader(file_path: str, password: str = None) -> list[Document]:
    """
    加载PDF文件内容
    :param file_path:PDF文件路径
    :param password:PDF密码（如果有）
    :return: PDF文件内容
    """
    # 处理路径，确保使用的是绝对路径
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    loader = PyPDFLoader(abs_file_path, password=password)
    return await asyncio.to_thread(loader.load)


async def txt_loader(file_path: str) -> list[Document]:
    """
    加载TXT文件内容
    :param file_path:TXT文件路径
    :return: TXT文件内容
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path

    # 使用不同的编码加载文件
    encodings = ['utf-8', 'gbk']
    for encoding in encodings:
        try:
            loader = TextLoader(abs_file_path, encoding=encoding)
            return await asyncio.to_thread(loader.load)
        except Exception as e:
            logger.error(f"[文本文件加载失败]使用编码{encoding}加载{abs_file_path}时出错:{e}")
            continue

    # 所有编码都失败，返回空列表
    return []


async def word_loader(file_path: str) -> list[Document]:
    """
    加载WORD文件（使用python-docx提取文本）
    :param file_path: WORD文档路径
    :return: 文档对象列表
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error(f"[WORD文件加载]python-docx库未安装，无法加载DOCX文件")
            return []
        doc = DocxDocument(abs_file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        if not text.strip():
            logger.warning(f"[WORD文件加载]{abs_file_path}内容为空")
            return []
        return [Document(page_content=text, metadata={"source": abs_file_path})]
    except Exception as e:
        logger.error(f"[WORD文件加载]加载{abs_file_path}时出错:{e}")
        return []


async def markdown_loader(file_path: str) -> list[Document]:
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredMarkdownLoader(abs_file_path)
        return await asyncio.to_thread(loader.load)
    except Exception as e:
        logger.error(f"[Markdown文件加载]加载文件{abs_file_path}时出错:{e}")
        return []


async def ppt_loader(file_path: str) -> list[Document]:
    """
    加载PPT/PPTX文件内容
    :param file_path:
    :return:
    """
    abs_file_path = get_abstract_path(file_path) if not os.path.isabs(file_path) else file_path
    try:
        loader = UnstructuredPowerPointLoader(abs_file_path)
        return await asyncio.to_thread(loader.load)
    except Exception as e:
        logger.error(f"[PPT文件加载]加载文件{abs_file_path}时出错:{e}")
        return []